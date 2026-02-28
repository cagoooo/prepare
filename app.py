import os
import io
from flask import Flask, render_template, request, jsonify, send_file
from google import genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from docx.enum.style import WD_STYLE_TYPE
import traceback
from linebot import LineBotApi
from linebot.models import TextSendMessage, FlexSendMessage


# 載入環境變數
load_dotenv()

app = Flask(__name__)
# 診斷路徑
import os
print(f"DEBUG: App file: {__file__}")
print(f"DEBUG: CWD: {os.getcwd()}")
print(f"DEBUG: Template folder: {app.template_folder}")

# 停用模板緩存
app.config['TEMPLATES_AUTO_RELOAD'] = True
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0

# 配置 Gemini API
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    print("Warning: GEMINI_API_KEY is not set. AI features will not work.")
    client = None
else:
    client = genai.Client(api_key=GEMINI_API_KEY)

@app.route('/')
def index():
    return render_template('index.html')

def html_to_docx(html_content):
    doc = Document()
    doc.add_heading('教學活動設計', 0)

    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    table = soup.find('table')

    if table:
        # Create a table in the Word document
        rows = table.find_all('tr')
        docx_table = doc.add_table(rows=len(rows), cols=2)
        docx_table.style = 'Table Grid'
        docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                docx_cell = docx_table.cell(i, j)

                if '教學活動內容及實施方式' in cell.get_text():
                    # Clear the cell content
                    docx_cell.text = ''

                    # Split the content by <br> tags
                    content_parts = cell.decode_contents().split('<br>')
                    current_section = None

                    for part in content_parts:
                        # Remove HTML tags
                        clean_part = BeautifulSoup(part, 'html.parser').get_text(strip=True)

                        if clean_part:
                            if any(heading in clean_part for heading in ['引起動機', '發展活動', '綜合活動']):
                                p = docx_cell.add_paragraph()
                                p.add_run(clean_part).bold = True
                                current_section = clean_part
                            else:
                                if current_section:
                                    p = docx_cell.add_paragraph(clean_part, style='List Bullet')
                                else:
                                    p = docx_cell.add_paragraph(clean_part)
                else:
                    docx_cell.text = cell.get_text(strip=True)

                # Apply formatting to cells
                for paragraph in docx_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Microsoft JhengHei'

                # Apply bold formatting to header cells
                if cell.name == 'th':
                    for paragraph in docx_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Add yellow highlight to cells containing "（僅供參考）"
                if "（僅供參考）" in docx_cell.text:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), "FFFF00")
                    docx_cell._tc.get_or_add_tcPr().append(shading_elm)

    # Ensure 'List Bullet' style exists in the document
    if 'List Bullet' not in doc.styles:
        doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        doc.styles['List Bullet'].base_style = doc.styles['Normal']
        doc.styles['List Bullet'].paragraph_format.left_indent = Inches(0.25)
        doc.styles['List Bullet'].paragraph_format.first_line_indent = Inches(-0.25)

    # Save the document to a BytesIO object
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)

    return docx_file

def html_to_email_friendly_table(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    email_content = '<table style="border-collapse: collapse; width: 100%;">'

    for row in soup.find_all('tr'):
        email_content += '<tr style="border: 1px solid #ddd;">'
        cells = row.find_all(['th', 'td'])
        if len(cells) == 2:
            label = cells[0].get_text(strip=True)
            content = cells[1].get_text(strip=True)

            # 特殊處理教學活動內容及實施方式
            if '教學活動內容及實施方式' in label:
                email_content += f'<td style="border: 1px solid #ddd; padding: 8px; font-weight: bold;">{label}</td>'
                email_content += '<td style="border: 1px solid #ddd; padding: 8px;">'
                parts = ['引起動機', '發展活動', '綜合活動']
                for part in parts:
                    if part in content:
                        try:
                            section = content.split(part, 1)[1]
                            next_part = next((p for p in parts if p in section), None)
                            if next_part:
                                section = section.split(next_part, 1)[0]
                            email_content += f'<strong>{part}:</strong><br>{section.strip()}<br><br>'
                        except IndexError:
                            # 如果分割失敗，直接使用整個內容
                            email_content += f'<strong>{part}:</strong><br>{content}<br><br>'
                email_content += '</td>'
            else:
                email_content += f'<td style="border: 1px solid #ddd; padding: 8px; font-weight: bold;">{label}</td>'
                email_content += f'<td style="border: 1px solid #ddd; padding: 8px;">{content}</td>'

        email_content += '</tr>'

    email_content += '</table>'
    return email_content

def create_lesson_plan_flex_message(data, html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    table = soup.find('table')
    
    body_contents = []
    
    if table:
        rows = table.find_all('tr')
        for row in rows:
            cells = row.find_all(['th', 'td'])
            if len(cells) >= 2:
                key = cells[0].get_text(strip=True).replace('（僅供參考）', '')
                # 替換 <br> 為換行符
                for br in cells[1].find_all('br'):
                    br.replace_with('\n')
                value = cells[1].get_text(strip=True)
                
                if not value:
                    continue
                
                # LINE 文字組件字數限制防禦
                if len(value) > 1900:
                    value = value[:1900] + "...\n(內容過長，請至網頁或 Word 檔查看完整內容)"
                
                body_contents.append({
                    "type": "box",
                    "layout": "vertical",
                    "margin": "lg",
                    "spacing": "sm",
                    "contents": [
                        {
                            "type": "text",
                            "text": f"📌 {key}",
                            "weight": "bold",
                            "color": "#1DB446",
                            "size": "sm",
                            "wrap": True
                        },
                        {
                            "type": "text",
                            "text": value,
                            "wrap": True,
                            "size": "sm",
                            "color": "#333333"
                        }
                    ]
                })
                body_contents.append({
                    "type": "separator",
                    "margin": "lg"
                })
        
        # 移除最後一個分隔線
        if body_contents and body_contents[-1]["type"] == "separator":
            body_contents.pop()

    flex_dict = {
      "type": "bubble",
      "size": "giga",
      "header": {
        "type": "box",
        "layout": "vertical",
        "backgroundColor": "#2C3E50",
        "paddingAll": "20px",
        "contents": [
          {
            "type": "text",
            "text": "🎓 教師數位備課小幫手",
            "color": "#FFFFFF",
            "weight": "bold",
            "size": "sm"
          },
          {
            "type": "text",
            "text": f"{data['subject']} - {data['unit']}",
            "color": "#F39C12",
            "weight": "bold",
            "size": "xl",
            "margin": "md",
            "wrap": True
          },
          {
            "type": "text",
            "text": f"適用年級：{data['grade']}",
            "color": "#BDC3C7",
            "size": "xs",
            "margin": "sm"
          }
        ]
      },
      "body": {
        "type": "box",
        "layout": "vertical",
        "backgroundColor": "#F8F9F9",
        "paddingAll": "20px",
        "contents": body_contents
      }
    }
    
    return flex_dict

@app.route('/download_docx', methods=['POST'])
def download_docx():
    try:
        html_content = request.json['html_content']
        docx_file = html_to_docx(html_content)

        # 添加 CSS 樣式來顯示表格框線
        styled_html_content = f"""
        <style>
            table {{ border-collapse: collapse; width: 100%; }}
            th, td {{ border: 1px solid black; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
        </style>
        {html_content}
        """

        # 建立郵件物件
        msg = Message('教案成功下載通知',
                      sender='210@mail2.smes.tyc.edu.tw',
                      recipients=['210@mail2.smes.tyc.edu.tw'])
        msg.body = "新的教案已生成，您可以在下方查看格式化的內容。"
        msg.html = styled_html_content

        # 發送郵件（使用獨立 try-except 隔離）
        try:
            mail.send(msg)
        except Exception as mail_err:
            print(f"Notification mail send failed: {str(mail_err)}")

        # 重置文件指針位置
        docx_file.seek(0)

        return send_file(
            docx_file,
            as_attachment=True,
            download_name='lesson_plan.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/generate_plan', methods=['POST'])
def generate_plan():
    if not GEMINI_API_KEY:
        return jsonify({"success": False, "error": "Gemini API key is not set"}), 500

    data = request.json
    # 移除可能導致編碼錯誤的 print
    # print(f"Received request data: {data}")

    prompt = f'''
請為以下教學活動生成一個完整的十二年國教教案，使用繁體中文及台灣常用詞彙：

教學領域名稱：{data['subject']}
實施年級：{data['grade']}
單元名稱：{data['unit']}

額外細節：{data['details']}

請生成一個包含以下欄位的 HTML 表格，標題為「<h3 class="lesson-plan-title">教學活動設計</h3>」：

1. 領域名稱
2. 設計者
3. 實施年級
4. 單元名稱
5. 總綱核心素養<span class="reference-only">（僅供參考）</span>
6. 領綱核心素養<span class="reference-only">（僅供參考）</span>
7. 核心素養呼應說明
8. 學習重點-學習表現<span class="reference-only">（僅供參考）</span>
9. 學習重點-學習內容<span class="reference-only">（僅供參考）</span>
10. 議題融入-實質內容<span class="reference-only">（僅供參考）</span>
11. 議題融入-所融入之學習重點
12. 教材來源
13. 教學資源-教師（請提供多樣化的資源，如教學影片、線上工具、實體教具等）
14. 教學資源-學生（請包含多元的學習材料，如閱讀文本、數位資源、實驗器材等）
15. 學習目標<span class="reference-only">（僅供參考）</span>
16. 教學重點（請提供3-5個具體且可衡量的重點）
17. 課前準備（包括教師和學生的準備事項）
18. 教學活動內容及實施方式
19. 課後延伸（提供多樣化的延伸活動建議，如專題研究、實地考察、創意作品等）
20. 評量方式（包括形成性評量和總結性評量，並提供多元評量方式）

請在生成的HTML表格中，確保在第5、6、8、9、10和15項的欄位名稱後面都加上帶有 'reference-only' 類的「（僅供參考）」標註。

請特別注意「教學活動內容及實施方式」部分，總時間為 40 分鐘，應包含以下詳細資訊：
1. 引起動機（約5分鐘）：
   - 描述如何吸引學生注意力並引導他們進入學習狀態
   - 提供創新的開場方式，如：情境模擬、問題探索、生活經驗連結等
2. 發展活動（約30分鐘）：
   a. 詳細說明每個教學步驟，包括每個步驟的時間分配
   b. 描述教師如何引導學生思考、討論或操作，運用多元教學策略如：
      - 合作學習
      - 探究式學習
      - 專題導向學習
      - 翻轉教學
      - 遊戲化學習
   c. 提供具體的問題示例、活動指引或討論主題
   d. 說明如何運用科技工具或多媒體資源增強學習體驗
   e. 描述差異化教學策略，以照顧不同程度的學生
3. 綜合活動（約5分鐘）：
   - 說明如何幫助學生整合所學知識
   - 提供創意的總結方式，如：概念圖、角色扮演、辯論等

請確保每個部分都有充分且具體的描述，包括時間分配，使教師能夠輕鬆理解並執行這個教案。請使用台灣教育常用的詞彙與表達方式，並融入適當的教育理念和創新教學方法。
'''

    # 移除可能導致編碼錯誤的 print
    # print(f"Generated prompt:\n{prompt}")

    try:
        print("Sending request to Gemini API...")
        response = client.models.generate_content(
            model="gemini-2.5-flash-lite",
            contents=prompt
        )
        content = response.text
        if not content:
            raise ValueError("Gemini returned an empty response.")

        # 移除 Markdown 標記
        content = content.replace('```html', '').replace('```', '').strip()
        if '</table>' in content:
            content = content.split('</table>')[0] + '</table>'

        # 將 HTML 內容轉換為郵件友好的表格格式
        try:
            email_friendly_content = html_to_email_friendly_table(content)
        except Exception as e:
            print(f"Error in converting HTML to email-friendly format: {str(e)}")
            email_friendly_content = content 
        
        # LINE 通知邏輯
        if line_bot_api and LINE_USER_ID:
            try:
                flex_content = create_lesson_plan_flex_message(data, content)
                line_bot_api.push_message(
                    LINE_USER_ID, 
                    FlexSendMessage(
                        alt_text=f"✨ 教案生成成功！({data['subject']} - {data['unit']})",
                        contents=flex_content
                    )
                )
                print("LINE Flex notification sent successfully.")
            except Exception as line_err:
                print(f"LINE notification failed: {str(line_err)}")
            
        return jsonify({"success": True, "plan": content, "html_content": content})
    except Exception as e:
        traceback.print_exc()
        print(f"Error occurred: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

# LINE API 配置
LINE_CHANNEL_ACCESS_TOKEN = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN")
LINE_USER_ID = os.environ.get("LINE_USER_ID")

if LINE_CHANNEL_ACCESS_TOKEN:
    line_bot_api = LineBotApi(LINE_CHANNEL_ACCESS_TOKEN)
else:
    print("Warning: LINE_CHANNEL_ACCESS_TOKEN is not set. LINE notifications will not work.")
    line_bot_api = None

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=False, threaded=True)
